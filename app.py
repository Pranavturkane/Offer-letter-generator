"""
go to line 23 24
Appointment Letter Generator - Flask App (Windows Compatible)
- No LibreOffice required
- Uses docx2pdf (calls MS Word on Windows automatically)
- Falls back to LibreOffice if Word not available
"""

from flask import Flask, request, send_file
import pandas as pd
import os
import re
import shutil
import zipfile
import traceback
from docx import Document
from pypdf import PdfReader, PdfWriter

app = Flask(__name__)

BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
#to run locally
# UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
# OUTPUT_FOLDER = os.path.join(BASE_DIR, "output")
#to run web
UPLOAD_FOLDER = "/tmp/uploads"
OUTPUT_FOLDER = "/tmp/output"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

DEFAULT_PASSWORD = "1234"

# ── Sample values in template -> Excel column names ───────────────────────────
REPLACEMENTS = [
    ("Sarwina VP",          "Full Name"),
    ("Sarwina",             "First Name"),
    ("Ms",                  "Title"),
    ("No.4,",               "Address 1"),
    ("4th street, Anna Nagar, Thirumullaivoyal", "Address 2"),
    ("Thirumullaivoyal, Chennai",                "Address 3"),
    ("600062",              "Pin Code"),
    ("6382260814",          "Mobile Number"),
    ("25-Feb-26",           "Letter Date"),
    ("04-Sep-26",           "End Date"),
    ("06-Mar-26",           "Start Date"),
    ("CEVA Freight Management Process", "Department"),
    ("Airoli",              "Location"),
    ("Six",                 "Internship Period"),
    ("32000",               "Stipend"),
    ("One Week",            "No. of Weeks/ Months Notice Period"),
    ("Ajay Mohan",          "Signing Authority"),
    ("Director - Human Resources", "Signing Designation"),
]


# ── Filename sanitizer ────────────────────────────────────────────────────────

def safe_filename(name: str) -> str:
    name = name.strip()
    name = re.sub(r'[\\/:*?"<>|]', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name or "Employee"


# ── Run merger + replacer ─────────────────────────────────────────────────────

def merge_runs(paragraph):
    if not paragraph.runs:
        return
    full_text = paragraph.text
    if not full_text.strip():
        return

    ref = paragraph.runs[0]
    for r in paragraph.runs:
        if r.text.strip():
            ref = r
            break

    p = paragraph._p
    for r in paragraph.runs:
        p.remove(r._r)

    new_run           = paragraph.add_run(full_text)
    new_run.bold      = ref.bold
    new_run.italic    = ref.italic
    new_run.underline = ref.underline
    if ref.font.size:
        new_run.font.size = ref.font.size
    if ref.font.name:
        new_run.font.name = ref.font.name
    try:
        if ref.font.color and ref.font.color.type:
            new_run.font.color.rgb = ref.font.color.rgb
    except Exception:
        pass


def replace_in_paragraph(paragraph, replacements):
    merge_runs(paragraph)
    for run in paragraph.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, str(new))


def fill_template(template_path, row_data):
    doc = Document(template_path)

    repl = {}
    for sample, col in REPLACEMENTS:
        val = row_data.get(col, sample)
        if val == "" or (not isinstance(val, str) and pd.isna(val)):
            val = sample
        repl[sample] = str(val).strip()

    for para in doc.paragraphs:
        replace_in_paragraph(para, repl)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, repl)
    return doc


# ── DOCX -> PDF (tries MS Word first, then LibreOffice) ──────────────────────

def docx_to_pdf(docx_path, out_dir):
    base     = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, base + ".pdf")

    # Method 1: docx2pdf (uses MS Word on Windows — best quality)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    # Method 2: LibreOffice (cross-platform fallback)
    try:
        import subprocess
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        lo_pdf = os.path.join(out_dir, base + ".pdf")
        if os.path.exists(lo_pdf):
            return lo_pdf
        raise RuntimeError(result.stderr.strip())
    except FileNotFoundError:
        pass

    raise RuntimeError(
        "Could not convert to PDF.\n"
        "Please install either:\n"
        "  - Microsoft Word (recommended for Windows)\n"
        "  - LibreOffice (free, cross-platform)\n"
        "Then run:  pip install docx2pdf"
    )


# ── PDF password protection ───────────────────────────────────────────────────

def encrypt_pdf(src, dst, password):
    reader = PdfReader(src)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password=password, owner_password=password)
    with open(dst, "wb") as f:
        writer.write(f)


# ── Main generation ───────────────────────────────────────────────────────────

def generate_documents(excel_path, template_path):
    if os.path.exists(OUTPUT_FOLDER):
        shutil.rmtree(OUTPUT_FOLDER)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    df = pd.read_excel(excel_path)
    df.columns = df.columns.str.strip()
    df = df.fillna("")

    pdf_files = []
    errors    = []

    for idx, row in df.iterrows():
        row_data  = row.to_dict()
        full_name = safe_filename(str(row_data.get("Full Name", f"Employee_{idx}")))
        pdf_name  = safe_filename(str(row_data.get("PdfFileName", full_name))) or full_name
        password  = str(row_data.get("Password", DEFAULT_PASSWORD)).strip() or DEFAULT_PASSWORD

        # Use index in temp name to avoid any spaces/special chars crashing conversion
        tmp_docx  = os.path.join(OUTPUT_FOLDER, f"tmp_{idx}.docx")
        raw_pdf   = os.path.join(OUTPUT_FOLDER, f"tmp_{idx}.pdf")
        final_pdf = os.path.join(OUTPUT_FOLDER, pdf_name + ".pdf")

        try:
            doc = fill_template(template_path, row_data)
            doc.save(tmp_docx)

            converted = docx_to_pdf(tmp_docx, OUTPUT_FOLDER)

            # Rename raw PDF to index-based name if needed
            if converted != raw_pdf and os.path.exists(converted):
                shutil.move(converted, raw_pdf)

            encrypt_pdf(raw_pdf, final_pdf, password)
            pdf_files.append(final_pdf)

        except Exception as e:
            errors.append(f"{full_name}:\n{traceback.format_exc()}")

        finally:
            for f in [tmp_docx, raw_pdf]:
                if os.path.exists(f):
                    try:
                        os.remove(f)
                    except Exception:
                        pass

    # Error log
    if errors:
        log_path = os.path.join(OUTPUT_FOLDER, "_ERRORS.txt")
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(errors))
        pdf_files.append(log_path)

    zip_path = os.path.join(OUTPUT_FOLDER, "All_Offer_Letters.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in pdf_files:
            zf.write(fp, os.path.basename(fp))

    return zip_path, len(pdf_files) - len(errors), errors


# ── Flask UI ──────────────────────────────────────────────────────────────────

HTML = """<!DOCTYPE html>
<html>
<head>
  <title>Offer Letter Generator</title>
  <style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: 'Segoe UI', Arial, sans-serif;
      background: #f0f2f5;
      display: flex; align-items: center; justify-content: center;
      min-height: 100vh;
    }}
    .card {{
      background: white; border-radius: 12px;
      box-shadow: 0 4px 24px rgba(0,0,0,0.10);
      padding: 40px 44px; width: 500px;
    }}
    h2 {{ color: #1a202c; font-size: 22px; margin-bottom: 6px; }}
    .subtitle {{ color: #718096; font-size: 13px; margin-bottom: 28px; }}
    label {{ display: block; font-weight: 600; color: #2d3748; margin-bottom: 6px; font-size: 14px; }}
    .field {{ margin-bottom: 20px; }}
    input[type=file] {{
      width: 100%; padding: 10px 12px;
      border: 2px dashed #cbd5e0; border-radius: 8px;
      background: #f7fafc; color: #4a5568; font-size: 13px; cursor: pointer;
    }}
    input[type=file]:hover {{ border-color: #4299e1; background: #ebf8ff; }}
    .hint {{ font-size: 11px; color: #a0aec0; margin-top: 5px; }}
    button {{
      width: 100%; padding: 13px;
      background: linear-gradient(135deg, #2b6cb0, #3182ce);
      color: white; border: none; border-radius: 8px;
      font-size: 15px; font-weight: 600; cursor: pointer;
      margin-top: 8px; transition: opacity .2s;
    }}
    button:hover {{ opacity: .88; }}
    button:disabled {{ opacity: .5; cursor: not-allowed; }}
    .error {{
      background: #fff5f5; border: 1px solid #feb2b2;
      color: #c53030; border-radius: 8px;
      padding: 12px 16px; margin-top: 18px; font-size: 13px;
      white-space: pre-wrap; word-break: break-word;
    }}
    .spinner {{ display:none; text-align:center; margin-top:14px; color:#718096; font-size:13px; }}
    .req {{ background:#fffbeb; border:1px solid #f6e05e; border-radius:8px; padding:12px 16px; margin-bottom:20px; font-size:12px; color:#744210; }}
  </style>
</head>
<body>
<div class="card">
  <h2>📄 Offer Letter Generator</h2>
  <p class="subtitle">Generates password-protected PDFs from your Excel data and Word template</p>

  <div class="req">
    <strong>Requirements:</strong> Make sure you have run:<br>
    <code>pip install flask pandas python-docx pypdf openpyxl docx2pdf</code><br>
    And have <strong>Microsoft Word</strong> installed (for PDF conversion).
  </div>

  <form method="POST" enctype="multipart/form-data"
        onsubmit="document.getElementById('spin').style.display='block';
                  this.querySelector('button').disabled=true;
                  this.querySelector('button').textContent='⏳ Generating...';">

    <div class="field">
      <label>📊 Excel File (.xlsx)</label>
      <input type="file" name="excel" accept=".xlsx,.xls" required>
    </div>

    <div class="field">
      <label>📝 Word Template (.docx)</label>
      <input type="file" name="template" accept=".docx" required>
      <p class="hint">Use the original Date.docx template (sample values like Sarwina VP, 25-Feb-26)</p>
    </div>

    <button type="submit">⚙️ Generate Offer Letters</button>
    <div class="spinner" id="spin">⏳ Please wait — converting letters to PDF…</div>
  </form>

  {error}
</div>
</body>
</html>"""


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        excel    = request.files.get("excel")
        template = request.files.get("template")

        if not excel or not template:
            return HTML.format(error='<div class="error">❌ Please upload both files.</div>')

        excel_path    = os.path.join(UPLOAD_FOLDER, safe_filename(excel.filename))
        template_path = os.path.join(UPLOAD_FOLDER, safe_filename(template.filename))
        excel.save(excel_path)
        template.save(template_path)

        try:
            zip_path, count, errors = generate_documents(excel_path, template_path)
            return send_file(zip_path, as_attachment=True,
                             download_name="All_Offer_Letters.zip",
                             mimetype="application/zip")
        except Exception as e:
            err_detail = traceback.format_exc()
            return HTML.format(
                error=f'<div class="error">❌ {str(e)}\n\n{err_detail}</div>'
            )

    return HTML.format(error="")


if __name__ == "__main__":
    print("=" * 55)
    print("  Offer Letter Generator")
    print("  Open in browser: http://127.0.0.1:5000")
    print("=" * 55)
    app.run(debug=False, host="127.0.0.1", port=5000, threaded=False)
